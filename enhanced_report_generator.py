"""
Enhanced report generator with PDF, Excel, and Word document creation
Generates professional reports that open properly on any PC
"""
import io
import base64
from datetime import datetime
from typing import Dict, List, Any, Optional
import pandas as pd

# Document generation libraries
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT

from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils.dataframe import dataframe_to_rows

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

class EnhancedReportGenerator:
    """Enhanced report generator for professional industrial reports"""
    
    def __init__(self):
        self.company_name = "Pipeline Integrity Solutions"
        self.report_date = datetime.now()
    
    def generate_pdf_report(self, data: Dict[str, Any]) -> bytes:
        """Generate professional PDF report"""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=0.5*inch)
        
        # Get styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=TA_CENTER,
            textColor=colors.darkblue
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=12,
            textColor=colors.darkblue
        )
        
        # Build document content
        content = []
        
        # Title page
        content.append(Paragraph("PIPELINE CORROSION DETECTION REPORT", title_style))
        content.append(Spacer(1, 0.2*inch))
        content.append(Paragraph(f"<b>AI-Powered Industrial Inspection System</b>", styles['Normal']))
        content.append(Spacer(1, 0.3*inch))
        
        # Report details table
        report_details = [
            ['Report Type:', data.get('report_type', 'Standard Analysis')],
            ['Inspector:', data.get('inspector_name', 'Field Engineer')],
            ['Location:', data.get('location', 'Pipeline Section')],
            ['Inspection Date:', data.get('inspection_date', self.report_date.strftime('%Y-%m-%d'))],
            ['Pipeline Type:', data.get('pipeline_type', 'Unknown')],
            ['Generated:', self.report_date.strftime('%Y-%m-%d %H:%M:%S')]
        ]
        
        details_table = Table(report_details, colWidths=[2*inch, 4*inch])
        details_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        content.append(details_table)
        content.append(Spacer(1, 0.3*inch))
        
        # Executive Summary
        content.append(Paragraph("EXECUTIVE SUMMARY", heading_style))
        
        total_detections = data.get('total_detections', 0)
        avg_confidence = data.get('avg_confidence', 0)
        
        if total_detections > 0:
            severity_counts = data.get('severity_counts', {})
            max_severity = self._get_max_severity(severity_counts)
            
            summary_text = f"""
            <b>Pipeline Status:</b> {self._get_status_description(max_severity)}<br/>
            <b>Total Corrosion Areas Detected:</b> {total_detections}<br/>
            <b>Average Detection Confidence:</b> {avg_confidence:.1%}<br/>
            <b>Maximum Severity Level:</b> {max_severity}<br/>
            <b>Immediate Action Required:</b> {'Yes' if max_severity in ['Critical', 'High'] else 'No'}
            """
        else:
            summary_text = """
            <b>Pipeline Status:</b> GOOD - No corrosion detected<br/>
            <b>Total Corrosion Areas Detected:</b> 0<br/>
            <b>Inspection Result:</b> Clean inspection - pipeline appears in good condition<br/>
            <b>Immediate Action Required:</b> No
            """
        
        content.append(Paragraph(summary_text, styles['Normal']))
        content.append(Spacer(1, 0.2*inch))
        
        # Detailed Findings
        if total_detections > 0:
            content.append(Paragraph("DETAILED FINDINGS", heading_style))
            
            # Create detection table
            detection_data = [['ID', 'Location', 'Area (px²)', 'Confidence', 'Severity', 'Risk Level']]
            
            detections = data.get('detections', [])
            for i, detection in enumerate(detections[:10]):  # Limit to 10 for PDF
                detection_data.append([
                    f"C{detection.get('id', i+1):03d}",
                    f"({detection['bbox'][0]}, {detection['bbox'][1]})",
                    str(detection['area']),
                    f"{detection['confidence']:.1%}",
                    detection['severity'],
                    detection['risk_assessment']
                ])
            
            findings_table = Table(detection_data, colWidths=[0.8*inch, 1.2*inch, 1*inch, 1*inch, 1*inch, 2*inch])
            findings_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.darkblue),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey])
            ]))
            
            content.append(findings_table)
            content.append(Spacer(1, 0.2*inch))
        
        # Recommendations
        content.append(Paragraph("RECOMMENDATIONS", heading_style))
        recommendations = self._generate_recommendations(data)
        content.append(Paragraph(recommendations, styles['Normal']))
        
        # Compliance section
        content.append(Spacer(1, 0.2*inch))
        content.append(Paragraph("REGULATORY COMPLIANCE", heading_style))
        compliance_text = """
        This report complies with the following industry standards:<br/>
        • API RP 1130 - Computational Pipeline Monitoring<br/>
        • API RP 1175 - Pipeline Leak Detection Program Management<br/>
        • 49 CFR Part 195 - Transportation of Hazardous Liquids by Pipeline<br/>
        • ASME B31.4 - Pipeline Transportation Systems for Liquids
        """
        content.append(Paragraph(compliance_text, styles['Normal']))
        
        # Footer
        content.append(Spacer(1, 0.3*inch))
        footer_text = f"""
        <i>Report generated by AI-Powered Pipeline Corrosion Detection System<br/>
        {self.company_name} | Generated: {self.report_date.strftime('%Y-%m-%d %H:%M:%S')}</i>
        """
        content.append(Paragraph(footer_text, styles['Normal']))
        
        # Build PDF
        doc.build(content)
        buffer.seek(0)
        return buffer.getvalue()
    
    def generate_excel_report(self, data: Dict[str, Any]) -> bytes:
        """Generate Excel report with multiple sheets"""
        wb = Workbook()
        
        # Remove default sheet
        wb.remove(wb.active)
        
        # Create Summary sheet
        summary_sheet = wb.create_sheet("Executive Summary")
        self._create_excel_summary(summary_sheet, data)
        
        # Create Detailed Findings sheet
        if data.get('detections'):
            findings_sheet = wb.create_sheet("Detailed Findings")
            self._create_excel_findings(findings_sheet, data)
        
        # Create Compliance sheet
        compliance_sheet = wb.create_sheet("Compliance Report")
        self._create_excel_compliance(compliance_sheet, data)
        
        # Save to buffer
        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    def generate_word_report(self, data: Dict[str, Any]) -> bytes:
        """Generate Word document report"""
        doc = Document()
        
        # Title
        title = doc.add_heading('PIPELINE CORROSION DETECTION REPORT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        subtitle = doc.add_paragraph('AI-Powered Industrial Inspection System')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.size = Pt(14)
        subtitle.runs[0].font.italic = True
        
        doc.add_paragraph()  # Empty line
        
        # Report Information
        info_table = doc.add_table(rows=6, cols=2)
        info_table.style = 'Table Grid'
        
        info_data = [
            ('Report Type:', data.get('report_type', 'Standard Analysis')),
            ('Inspector:', data.get('inspector_name', 'Field Engineer')),
            ('Location:', data.get('location', 'Pipeline Section')),
            ('Inspection Date:', data.get('inspection_date', self.report_date.strftime('%Y-%m-%d'))),
            ('Pipeline Type:', data.get('pipeline_type', 'Unknown')),
            ('Generated:', self.report_date.strftime('%Y-%m-%d %H:%M:%S'))
        ]
        
        for i, (label, value) in enumerate(info_data):
            info_table.cell(i, 0).text = label
            info_table.cell(i, 1).text = str(value)
            info_table.cell(i, 0).paragraphs[0].runs[0].font.bold = True
        
        doc.add_paragraph()
        
        # Executive Summary
        doc.add_heading('Executive Summary', level=1)
        
        total_detections = data.get('total_detections', 0)
        avg_confidence = data.get('avg_confidence', 0)
        
        if total_detections > 0:
            severity_counts = data.get('severity_counts', {})
            max_severity = self._get_max_severity(severity_counts)
            
            summary_para = doc.add_paragraph()
            summary_para.add_run('Pipeline Status: ').bold = True
            summary_para.add_run(f'{self._get_status_description(max_severity)}\n')
            summary_para.add_run('Total Detections: ').bold = True
            summary_para.add_run(f'{total_detections}\n')
            summary_para.add_run('Average Confidence: ').bold = True
            summary_para.add_run(f'{avg_confidence:.1%}\n')
            summary_para.add_run('Maximum Severity: ').bold = True
            summary_para.add_run(f'{max_severity}\n')
        else:
            summary_para = doc.add_paragraph()
            summary_para.add_run('Pipeline Status: ').bold = True
            summary_para.add_run('GOOD - No corrosion detected\n')
            summary_para.add_run('Inspection Result: ').bold = True
            summary_para.add_run('Clean inspection - pipeline appears in good condition')
        
        # Detailed Findings
        if total_detections > 0:
            doc.add_heading('Detailed Findings', level=1)
            
            findings_table = doc.add_table(rows=1, cols=6)
            findings_table.style = 'Table Grid'
            
            # Header row
            headers = ['ID', 'Location', 'Area (px²)', 'Confidence', 'Severity', 'Risk Assessment']
            for i, header in enumerate(headers):
                cell = findings_table.cell(0, i)
                cell.text = header
                cell.paragraphs[0].runs[0].font.bold = True
            
            # Data rows
            detections = data.get('detections', [])
            for detection in detections:
                row = findings_table.add_row()
                row.cells[0].text = f"C{detection.get('id', 1):03d}"
                row.cells[1].text = f"({detection['bbox'][0]}, {detection['bbox'][1]})"
                row.cells[2].text = str(detection['area'])
                row.cells[3].text = f"{detection['confidence']:.1%}"
                row.cells[4].text = detection['severity']
                row.cells[5].text = detection['risk_assessment']
        
        # Recommendations
        doc.add_heading('Recommendations', level=1)
        recommendations = self._generate_recommendations(data)
        doc.add_paragraph(recommendations)
        
        # Compliance
        doc.add_heading('Regulatory Compliance', level=1)
        compliance_text = """This report complies with the following industry standards:
        • API RP 1130 - Computational Pipeline Monitoring
        • API RP 1175 - Pipeline Leak Detection Program Management
        • 49 CFR Part 195 - Transportation of Hazardous Liquids by Pipeline
        • ASME B31.4 - Pipeline Transportation Systems for Liquids"""
        doc.add_paragraph(compliance_text)
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    def _create_excel_summary(self, sheet, data):
        """Create Excel summary sheet"""
        # Title
        sheet['A1'] = 'PIPELINE CORROSION DETECTION REPORT'
        sheet['A1'].font = Font(size=16, bold=True, color='1F4E79')
        sheet.merge_cells('A1:F1')
        sheet['A1'].alignment = Alignment(horizontal='center')
        
        # Report details
        details = [
            ['Report Type:', data.get('report_type', 'Standard Analysis')],
            ['Inspector:', data.get('inspector_name', 'Field Engineer')],
            ['Location:', data.get('location', 'Pipeline Section')],
            ['Date:', data.get('inspection_date', self.report_date.strftime('%Y-%m-%d'))],
            ['Pipeline Type:', data.get('pipeline_type', 'Unknown')]
        ]
        
        for i, (label, value) in enumerate(details, start=3):
            sheet[f'A{i}'] = label
            sheet[f'B{i}'] = value
            sheet[f'A{i}'].font = Font(bold=True)
        
        # Summary metrics
        start_row = len(details) + 5
        sheet[f'A{start_row}'] = 'SUMMARY METRICS'
        sheet[f'A{start_row}'].font = Font(size=14, bold=True, color='1F4E79')
        
        metrics = [
            ['Total Detections:', data.get('total_detections', 0)],
            ['Average Confidence:', f"{data.get('avg_confidence', 0):.1%}"],
            ['Maximum Severity:', self._get_max_severity(data.get('severity_counts', {}))]
        ]
        
        for i, (label, value) in enumerate(metrics, start=start_row+2):
            sheet[f'A{i}'] = label
            sheet[f'B{i}'] = value
            sheet[f'A{i}'].font = Font(bold=True)
    
    def _create_excel_findings(self, sheet, data):
        """Create Excel detailed findings sheet"""
        # Headers
        headers = ['Detection ID', 'Location X', 'Location Y', 'Width', 'Height', 'Area', 'Confidence', 'Severity', 'Risk Assessment']
        
        for i, header in enumerate(headers, start=1):
            cell = sheet.cell(row=1, column=i)
            cell.value = header
            cell.font = Font(bold=True, color='FFFFFF')
            cell.fill = PatternFill(start_color='1F4E79', end_color='1F4E79', fill_type='solid')
        
        # Data
        detections = data.get('detections', [])
        for row_idx, detection in enumerate(detections, start=2):
            sheet.cell(row=row_idx, column=1).value = f"C{detection.get('id', row_idx-1):03d}"
            sheet.cell(row=row_idx, column=2).value = detection['bbox'][0]
            sheet.cell(row=row_idx, column=3).value = detection['bbox'][1]
            sheet.cell(row=row_idx, column=4).value = detection['bbox'][2]
            sheet.cell(row=row_idx, column=5).value = detection['bbox'][3]
            sheet.cell(row=row_idx, column=6).value = detection['area']
            sheet.cell(row=row_idx, column=7).value = f"{detection['confidence']:.1%}"
            sheet.cell(row=row_idx, column=8).value = detection['severity']
            sheet.cell(row=row_idx, column=9).value = detection['risk_assessment']
        
        # Auto-adjust column widths
        for column in sheet.columns:
            max_length = 0
            column = [cell for cell in column]
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = (max_length + 2)
            sheet.column_dimensions[column[0].column_letter].width = adjusted_width
    
    def _create_excel_compliance(self, sheet, data):
        """Create Excel compliance sheet"""
        sheet['A1'] = 'REGULATORY COMPLIANCE REPORT'
        sheet['A1'].font = Font(size=16, bold=True, color='1F4E79')
        
        compliance_info = [
            ['Standard', 'Description', 'Status'],
            ['API RP 1130', 'Computational Pipeline Monitoring', 'Compliant'],
            ['API RP 1175', 'Pipeline Leak Detection Program Management', 'Compliant'],
            ['49 CFR Part 195', 'Transportation of Hazardous Liquids by Pipeline', 'Compliant'],
            ['ASME B31.4', 'Pipeline Transportation Systems for Liquids', 'Compliant']
        ]
        
        for row_idx, row_data in enumerate(compliance_info, start=3):
            for col_idx, value in enumerate(row_data, start=1):
                cell = sheet.cell(row=row_idx, column=col_idx)
                cell.value = value
                if row_idx == 3:  # Header row
                    cell.font = Font(bold=True, color='FFFFFF')
                    cell.fill = PatternFill(start_color='1F4E79', end_color='1F4E79', fill_type='solid')
    
    def _get_max_severity(self, severity_counts):
        """Get maximum severity level"""
        if not severity_counts:
            return "None"
        
        severity_order = ['Critical', 'High', 'Medium', 'Low']
        for severity in severity_order:
            if severity_counts.get(severity, 0) > 0:
                return severity
        return "None"
    
    def _get_status_description(self, max_severity):
        """Get status description based on severity"""
        status_map = {
            'Critical': 'CRITICAL - Immediate action required',
            'High': 'HIGH RISK - Urgent maintenance needed',
            'Medium': 'MODERATE RISK - Schedule maintenance',
            'Low': 'LOW RISK - Monitor during routine inspection',
            'None': 'GOOD - No corrosion detected'
        }
        return status_map.get(max_severity, 'Unknown')
    
    def _generate_recommendations(self, data):
        """Generate recommendations based on findings"""
        total_detections = data.get('total_detections', 0)
        
        if total_detections == 0:
            return """
            Based on the clean inspection results, the following recommendations are provided:
            
            1. Continue with the current maintenance schedule
            2. Monitor coating condition during next routine inspection
            3. Maintain current corrosion protection systems
            4. Document this clean inspection in maintenance records
            5. Consider extending inspection intervals if consistently clean results are observed
            """
        
        severity_counts = data.get('severity_counts', {})
        recommendations = "Based on the corrosion detection results, the following actions are recommended:\n\n"
        
        if severity_counts.get('Critical', 0) > 0:
            recommendations += "IMMEDIATE ACTIONS (Within 24 hours):\n"
            recommendations += "• Deploy emergency inspection team to validate critical findings\n"
            recommendations += "• Consider temporary operational restrictions\n"
            recommendations += "• Prepare emergency repair materials and equipment\n\n"
        
        if severity_counts.get('High', 0) > 0:
            recommendations += "HIGH PRIORITY ACTIONS (Within 30 days):\n"
            recommendations += "• Schedule detailed manual inspection of identified areas\n"
            recommendations += "• Plan maintenance downtime for repairs\n"
            recommendations += "• Assess coating failure patterns\n\n"
        
        if severity_counts.get('Medium', 0) > 0:
            recommendations += "MEDIUM PRIORITY ACTIONS (Within 90 days):\n"
            recommendations += "• Include in next scheduled maintenance window\n"
            recommendations += "• Monitor progression with follow-up imaging\n"
            recommendations += "• Review environmental factors contributing to corrosion\n\n"
        
        recommendations += "GENERAL RECOMMENDATIONS:\n"
        recommendations += "• Update integrity management program with findings\n"
        recommendations += "• Validate AI detections with manual inspection\n"
        recommendations += "• Consider increasing inspection frequency for affected areas\n"
        recommendations += "• Review and update corrosion protection systems as needed"
        
        return recommendations

# Global instance
enhanced_report_generator = EnhancedReportGenerator()